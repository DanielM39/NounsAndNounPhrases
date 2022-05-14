#! python3.9
from textblob import TextBlob
from sys import argv, platform
import docx

def getTextDocx(filename):
	doc = docx.Document(filename)
	fullText = []
	for para in doc.paragraphs:
		fullText.append(para.text)
	for table in doc.tables:
		for row in table.rows:
			for cel in row.cells:
				fullText.append(cel.text)
	return '\n'.join(fullText)

def getTextDoc(filename):
	import win32com.client

	word = win32com.client.Dispatch("Word.Application")
	word.visible = False
	wb = word.Documents.Open(filename)
	doc = word.ActiveDocument
	txt = doc.Range().Text
	wb.Close()
	return txt

def getTextTxt(filename):
	txt = ""
	with open(filename, 'r', encoding='utf8') as f:
			for line in f:
				txt = txt + ' ' + line
	return txt

EXCLUTIONS = ["preconditions", "post-conditions", "scenario", "actor", "yyyy/mm/dd", "extensions", "merges", "step",
	"case document name", "alternative path", "ui actions", "primary actor", "customer stakeholders", "main success",
	"basic flow", "business layer", "system response", "case diagram", "main scenario", "alternative flows", 
	"alternative path", "scenario branches"]

def remove_duplicates(words):
	return list(dict.fromkeys(words))

if len(argv) > 1:
	txt = ""
	if argv[1].lower().endswith(".docx"):
		txt = getTextDocx(argv[1])
	elif argv[1].lower().endswith(".doc") and platform == "win32":
		txt = getTextDoc(argv[1])
	else:
		txt = getTextTxt(argv[1])
	blob = TextBlob(txt)
	
	with open("{}csv".format(argv[1].replace(argv[1].split('.')[-1],'')), 'w', encoding="utf8") as csv:
		csv.write("Noun,Meaning,Class?,Reason\n")
		for noun in remove_duplicates(blob.noun_phrases):
			if (noun.strip().lower() not in EXCLUTIONS) and (not noun.strip().isnumeric()):
				csv.write("{}\n".format(noun.title()))
else:
	print("Usage:\n\t %s input\nSupports: .docx, .doc, .txt" % argv[0].split('\\')[-1])